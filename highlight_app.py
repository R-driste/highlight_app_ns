import tkinter as tk
from tkinter import filedialog, messagebox
import os
from docx import Document
from itertools import groupby

# Main window
root = tk.Tk()
root.title("Transcription Highlight Widget")
root.geometry("400x400")

labela = tk.Label(root, text="File 1", font=("Arial", 12))
labela.pack(pady=20)
label1 = tk.Label(root, text="NONE SELECTED", font=("Arial", 12))
label1.pack(pady=20)
labelb = tk.Label(root, text="File 2", font=("Arial", 12))
labelb.pack(pady=20)
label2 = tk.Label(root, text="NONE SELECTED", font=("Arial", 12))
label2.pack(pady=10)

out_num = 0
file_names = {'file1': None, 'file2': None}
with open("file_num.txt", "r") as f:
    out_num = int(f.read())

def open_file_picker(file_num, label):
    try:
        file_path = filedialog.askopenfilename(filetypes=[("All Files", "*.*")])
        if file_path:
            if file_path.lower().endswith(('.doc', '.docx')):
                label.config(text=f"File: {file_path}")
                file_names[file_num] = file_path
            else:
                messagebox.showerror("Invalid File", "Please select a Microsoft Word document (.docx or .doc).")
                label.config(text="NONE SELECTED")
        else:
            label.config(text="NONE SELECTED")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred while selecting the file: {e}")
        label.config(text="NONE SELECTED")

    print(file_names)

output_folder_path = tk.StringVar()

def compare_time(out_num):
    #Check that the file is valid
    if file_names['file1'] is None or file_names['file2'] is None:
        messagebox.showerror("Error", "Please select both files before comparing.")
        return
    folder_path = filedialog.askdirectory(title="Select Output Folder")
    if folder_path:
        output_folder_path.set(folder_path)
    output_file_name = f"COMBINE_HIGHLIGHT_{out_num}.docx"
    output_file_path = os.path.join(output_folder_path.get(), output_file_name)

    #Create the input/output word documents
    file1 = Document(file_names['file1'])
    file2 = Document(file_names['file2'])
    outfile = Document()
    outfile._body.clear_content()

    #Loop through the paragraphs and compare them
    for para1, para2 in zip(file1.paragraphs, file2.paragraphs):
        runs1 = [run.text for run in para1.runs]
        runs2 = [run.text for run in para2.runs]
        final_text = ''.join(runs1)
        final_map1 = ""
        final_map2 = ""
        
        for run in para1.runs:
            if run.font.highlight_color == None:
                final_map1 += "N" * len(run.text)
            else:
                final_map1 += "Y" * len(run.text)

        for run in para2.runs:
            if run.font.highlight_color == None:
                final_map2 += "N" * len(run.text)
            else:
                final_map2 += "Y" * len(run.text)
        map_FINAL = ""

        #Calculate overlaps
        for char1, char2 in zip(final_map1, final_map2):
            if char1 == char2:
                if char1=="Y":
                    map_FINAL += "G"
                elif char1=="N":
                    map_FINAL += "N"
            else:
                map_FINAL += "Y"
        
        final_group = [(char, len(list(group))) for char, group in groupby(map_FINAL)]

        #Now that we know the highlighting, create the output paragraph

        p = outfile.add_paragraph('')
        

        for group in final_group:
            r = p.add_run(final_text[:group[1]])
            final_text = final_text[group[1]:]
            if group[0] == "N":
                r.font.highlight_color = None
            elif group[0] == "Y":
                r.font.highlight_color = 7  # Yellow
            elif group[0] == "G":
                r.font.highlight_color = 4  # Green

    #Save the output file
    outfile.save(output_file_path)
    outfile = Document(output_file_path)
    for para in outfile.paragraphs:
        print(para.text)
        
    messagebox.showinfo("Success", f"Comparison saved to {output_file_path}")
    out_num += 1
    with open("file_num.txt", "w") as f:
        f.write(str(out_num))

#Decision Buttons
button1 = tk.Button(root, text="Select File 1", command=lambda: open_file_picker('file1', label1))
button1.pack(pady=10)
button2 = tk.Button(root, text="Select File 2", command=lambda: open_file_picker('file2', label2))
button2.pack(pady=10)
button3 = tk.Button(root, text="Create Comparison", command=lambda: compare_time(out_num))
button3.pack(pady=10)

#Run
root.mainloop()