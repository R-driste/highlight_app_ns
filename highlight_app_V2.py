import tkinter as tk
from tkinter import filedialog, messagebox
import os
from docx import Document
from itertools import groupby
from collections import Counter

#App Folder Selection Component
root = tk.Tk()
root.title("Transcription Highlight Widget")
root.geometry("450x600")

label0 = tk.Label(root, text="Welcome. Please select correct folders.\nFiles must be MMDDYYYY-#-IN Microsoft Doc format.\n\nKEY:\nEmpty: No highlight\nYellow: One highlight\nGreen: Two highlights\nBlue: Three highlights.", font=("Arial", 14))
label0.pack(pady=20)

labela = tk.Label(root, text="Input Folder:", font=("Arial", 15))
labela.pack(pady=20)
label1 = tk.Label(root, text="NONE SELECTED", font=("Arial", 10))
label1.pack(pady=20)
button1 = tk.Button(root, text="Select Input Folder", command=lambda: open_folder_picker('folder1', label1))
button1.pack(pady=10)

labelb = tk.Label(root, text="Output Folder:", font=("Arial", 15))
labelb.pack(pady=20)
label2 = tk.Label(root, text="NONE SELECTED", font=("Arial", 10))
label2.pack(pady=10)
button2 = tk.Button(root, text="Select Output Folder", command=lambda: open_folder_picker('folder2', label2))
button2.pack(pady=10)

folder_paths = {'folder1': None, 'folder2': None}

def open_folder_picker(folder_num, label):
    try:
        folder_path = filedialog.askdirectory(title="Select Folder")
        if folder_path:
            label.config(text=f"Folder: {folder_path}")
            folder_paths[folder_num] = folder_path
        else:
            label.config(text="NONE SELECTED")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred while selecting the folder: {e}")
        label.config(text="NONE SELECTED")
    output_folder_path = tk.StringVar()

def compare_time():
    #Grab All Files
    if folder_paths['folder1'] is None or folder_paths['folder2'] is None:
        messagebox.showerror("Error", "Please select both folders before comparing.")
        return
    file_reads = {}
    input = folder_paths['folder1']
    for file in os.listdir(input):
        if file.lower().endswith(('.doc', '.docx')):
            details = file.split("-")
            key = details[0] + "-" + details[1]
            if key not in file_reads.keys():
                file_reads[key] = []
            file_reads[key].append(file)
    print(file_reads)

    for key, values in file_reads.items():
        output_path = os.path.join(folder_paths['folder2'], f"COMBINED_{key}.docx")
        full_paths = [os.path.join(folder_paths['folder1'], value) for value in values]  # Ensure full paths
        if len(full_paths) == 2:
            print("Comparing 2 files")
            compare_2(full_paths, output_path)
        elif len(full_paths) == 3:
            print("Comparing 3 files")
            compare_3(full_paths, output_path)
        else:
            print(key, "needs to have 2 or 3 values")
    
    messagebox.showinfo("Success", "All files have been successfully processed!")

def compare_2(values, output_path):
    file1 = Document(values[0])
    file2 = Document(values[1])
    outfile = Document()
    outfile._body.clear_content()

    # Loop through the paragraphs and compare them
    for para1, para2 in zip(file1.paragraphs, file2.paragraphs):
        runs1 = [run.text for run in para1.runs]
        final_text = ''.join(runs1)
        final_map1 = ""
        final_map2 = ""

        for run in para1.runs:
            print(f"Run text: {run.text}, Highlight: {run.font.highlight_color}")  # Debug print
            if run.font.highlight_color is None:
                final_map1 += "N" * len(run.text)
            else:
                final_map1 += "Y" * len(run.text)

        for run in para2.runs:
            print(f"Run text: {run.text}, Highlight: {run.font.highlight_color}")  # Debug print
            if run.font.highlight_color is None:
                final_map2 += "N" * len(run.text)
            else:
                final_map2 += "Y" * len(run.text)

        map_FINAL = ""

        for char1, char2 in zip(final_map1, final_map2):
            if char1 == char2:
                if char1=="Y":
                    map_FINAL += "G"
                elif char1=="N":
                    map_FINAL += "N"
            else:
                map_FINAL += "Y"
        
        final_group = [(char, len(list(group))) for char, group in groupby(map_FINAL)]

        #Now that we know the highlighting, create the output paragraph
        p = outfile.add_paragraph('')
        for group in final_group:
            r = p.add_run(final_text[:group[1]])
            final_text = final_text[group[1]:]
            if group[0] == "N":
                r.font.highlight_color = None
            elif group[0] == "Y":
                r.font.highlight_color = 7  # Yellow
            elif group[0] == "G":
                r.font.highlight_color = 4  # Green

    #Save the output file
    outfile.save(output_path)

def compare_3(values, output_path):
    print("VALS:", values[0], values[1], values[2])
    file1 = Document(values[0])
    file2 = Document(values[1])
    file3 = Document(values[2])
    outfile = Document()
    outfile._body.clear_content()

    #Loop through the paragraphs and compare them
    for para1, para2, para3 in zip(file1.paragraphs, file2.paragraphs, file3.paragraphs):
        runs1 = [run.text for run in para1.runs]
        final_text = ''.join(runs1)
        final_map1 = ""
        final_map2 = ""
        final_map3 = ""
        
        for run in para1.runs:
            if run.font.highlight_color == None:
                final_map1 += "N" * len(run.text)
            else:
                final_map1 += "Y" * len(run.text)

        for run in para2.runs:
            if run.font.highlight_color == None:
                final_map2 += "N" * len(run.text)
            else:
                final_map2 += "Y" * len(run.text)
        
        for run in para3.runs:
            print(run.font.highlight_color)
            if run.font.highlight_color == None:
                final_map3 += "N" * len(run.text)
            else:
                final_map3 += "Y" * len(run.text)
        #print("RUNS")
        #print("RUN",final_map1, "\nRUN", final_map2,  "\nRUN", final_map3)
        map_FINAL = ""

        #Calculate overlaps
        for char1, char2, char3 in zip(final_map1, final_map2, final_map3):
            combined = char1 + char2 + char3
            counts = Counter(combined)
            if counts['Y'] == 0:
                map_FINAL += "N"
            elif counts['Y'] == 1:
                map_FINAL += "Y"
            elif counts['Y'] == 2:
                map_FINAL += "G"
            elif counts['Y'] == 3:
                map_FINAL += "B"
                
        final_group = [(char, len(list(group))) for char, group in groupby(map_FINAL)]
        #print("---")
        #print(map_FINAL)

        #Now that we know the highlighting, create the output paragraph

        p = outfile.add_paragraph('')
        for group in final_group:
            r = p.add_run(final_text[:group[1]])
            final_text = final_text[group[1]:]
            if group[0] == "N":
                r.font.highlight_color = None
            elif group[0] == "Y":
                r.font.highlight_color = 7  # Yellow
            elif group[0] == "G":
                r.font.highlight_color = 4
            elif group[0] == "B":
                r.font.highlight_color = 5

    #Save the output file
    outfile.save(output_path)
        
button3 = tk.Button(root, text="Create Comparison Highlights", command=lambda: compare_time(), font=("Arial", 16), width=25, height=2)
button3.pack(pady=20)  # Increase padding

#Run
root.mainloop()